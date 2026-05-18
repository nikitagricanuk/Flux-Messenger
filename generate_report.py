#!/usr/bin/env python3
"""Generate Flux Messenger coursework report (DOCX) per STO-005."""

import io
import os
import textwrap
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENTATION

# ─── STO-005 constants ────────────────────────────────────────────────────────
FONT_NAME   = "Times New Roman"
FONT_SIZE   = Pt(14)
FONT_SIZE_SM = Pt(12)
LEFT_MARGIN  = Cm(3.0)
RIGHT_MARGIN = Cm(1.0)
TOP_MARGIN   = Cm(1.5)
BOT_MARGIN   = Cm(2.0)
PARA_INDENT  = Cm(1.25)

# ─── Helper: set paragraph font ───────────────────────────────────────────────

def _set_run_font(run, bold=False, size=None):
    run.font.name  = FONT_NAME
    run.font.size  = size or FONT_SIZE
    run.font.bold  = bold
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    FONT_NAME)
    rFonts.set(qn('w:hAnsi'),    FONT_NAME)
    rFonts.set(qn('w:cs'),       FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rPr.insert(0, rFonts)


def _fmt_para(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, space_before=0, space_after=0):
    pf = para.paragraph_format
    pf.alignment    = align
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if indent:
        pf.first_line_indent = PARA_INDENT


def add_para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             indent=True, space_before=0, space_after=0, size=None):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    _set_run_font(run, bold=bold, size=size)
    _fmt_para(para, align=align, indent=indent,
              space_before=space_before, space_after=space_after)
    return para


def add_heading(doc, text, level, numbered=None):
    """level 1 = раздел, level 2 = подраздел"""
    para  = doc.add_paragraph()
    label = (f"{numbered} " if numbered else "") + text
    run   = para.add_run(label)
    _set_run_font(run, bold=True)
    pf = para.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = PARA_INDENT
    pf.space_before      = Pt(12)
    pf.space_after       = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # keep with next
    pPr = para._p.get_or_add_pPr()
    kn = OxmlElement('w:keepNext')
    pPr.append(kn)
    return para


def add_page_break(doc):
    para = doc.add_paragraph()
    run  = para.add_run()
    run.add_break(docx_break_type())
    return para


def docx_break_type():
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _el
    br = _el('w:br')
    br.set(_qn('w:type'), 'page')
    return br   # will be added via run.add_break workaround


def page_break(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run  = para.add_run()
    from docx.oxml import OxmlElement
    br   = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    return para


def add_figure_caption(doc, number, title):
    text = f"Рисунок {number} – {title}"
    para = doc.add_paragraph()
    run  = para.add_run(text)
    _set_run_font(run, bold=False)
    pf = para.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Pt(0)
    pf.space_before      = Pt(3)
    pf.space_after       = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return para


def add_table_caption(doc, number, title):
    text = f"Таблица {number} – {title}"
    para = doc.add_paragraph()
    run  = para.add_run(text)
    _set_run_font(run, bold=False)
    pf = para.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Pt(0)
    pf.space_before      = Pt(6)
    pf.space_after       = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return para


def embed_image(doc, img_bytes, width_cm=14.0):
    para = doc.add_paragraph()
    pf   = para.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Pt(0)
    pf.space_before      = Pt(6)
    pf.space_after       = Pt(3)
    run  = para.add_run()
    run.add_picture(img_bytes, width=Cm(width_cm))
    return para


def add_spec_table(doc, caption, fields, methods):
    """Add specification table with fields and methods."""
    if fields:
        add_table_caption(doc, caption + "а", "Поля класса")
        t = doc.add_table(rows=1, cols=4)
        t.style = 'Table Grid'
        hdr = t.rows[0].cells
        for i, h in enumerate(["№", "Поле", "Тип", "Описание"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                _set_run_font(run, bold=True, size=Pt(12))
        for idx, (name, typ, desc) in enumerate(fields, 1):
            row = t.add_row().cells
            row[0].text = str(idx)
            row[1].text = name
            row[2].text = typ
            row[3].text = desc
            for cell in row:
                for run in cell.paragraphs[0].runs:
                    _set_run_font(run, size=Pt(12))
        doc.add_paragraph()

    if methods:
        add_table_caption(doc, caption + "б", "Методы класса")
        t = doc.add_table(rows=1, cols=5)
        t.style = 'Table Grid'
        hdr = t.rows[0].cells
        for i, h in enumerate(["№", "Метод", "Тип", "Доступ", "Описание"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                _set_run_font(run, bold=True, size=Pt(12))
        for idx, (name, ret, acc, desc) in enumerate(methods, 1):
            row = t.add_row().cells
            row[0].text = str(idx)
            row[1].text = name
            row[2].text = ret
            row[3].text = acc
            row[4].text = desc
            for cell in row:
                for run in cell.paragraphs[0].runs:
                    _set_run_font(run, size=Pt(12))
        doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
#  DIAGRAM GENERATORS
# ══════════════════════════════════════════════════════════════════════════════

def _buf():
    return io.BytesIO()


def make_usecase_diagram():
    fig, ax = plt.subplots(figsize=(14, 9))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.set_facecolor('white')
    fig.patch.set_facecolor('white')

    # System boundary
    rect = mpatches.FancyBboxPatch((2.5, 0.5), 9.5, 8.0, boxstyle="round,pad=0.1",
                                    linewidth=2, edgecolor='black', facecolor='#f0f8ff', zorder=1)
    ax.add_patch(rect)
    ax.text(7.25, 8.65, 'Система аутентификации Flux Messenger',
            ha='center', va='center', fontsize=11, fontweight='bold')

    # Actor
    def draw_actor(ax, x, y, label):
        circle = plt.Circle((x, y + 0.4), 0.22, color='black', fill=False, linewidth=1.5, zorder=5)
        ax.add_patch(circle)
        ax.plot([x, x], [y + 0.18, y - 0.3], 'k-', linewidth=1.5, zorder=5)
        ax.plot([x - 0.3, x + 0.3], [y, y], 'k-', linewidth=1.5, zorder=5)
        ax.plot([x, x - 0.25], [y - 0.3, y - 0.65], 'k-', linewidth=1.5, zorder=5)
        ax.plot([x, x + 0.25], [y - 0.3, y - 0.65], 'k-', linewidth=1.5, zorder=5)
        ax.text(x, y - 0.9, label, ha='center', va='center', fontsize=9, fontweight='bold',
                wrap=True, multialignment='center')

    draw_actor(ax, 1.2, 5.5, 'Пользователь')

    # Use cases (ellipses)
    use_cases = [
        (6.2, 7.5, 'Войти по паролю'),
        (6.2, 6.2, 'Зарегистрироваться\nпо паролю'),
        (6.2, 4.9, 'Войти с Passkey'),
        (6.2, 3.6, 'Зарегистрировать\nPasskey'),
        (10.5, 7.5, 'Валидация данных\nформы'),
        (10.5, 6.0, 'Сохранение\nтокенов'),
        (10.5, 4.5, 'Обновление\nтокена доступа'),
        (10.5, 3.0, 'Биометрическая\nаутентификация'),
        (6.2, 2.2, 'Перейти к\nглавному экрану'),
    ]

    for (x, y, label) in use_cases:
        ellipse = mpatches.Ellipse((x, y), 3.0, 0.8, linewidth=1.5,
                                    edgecolor='navy', facecolor='white', zorder=3)
        ax.add_patch(ellipse)
        ax.text(x, y, label, ha='center', va='center', fontsize=8, zorder=4,
                multialignment='center')

    # Connections actor -> use cases
    uc_main = [(6.2, 7.5), (6.2, 6.2), (6.2, 4.9), (6.2, 3.6), (6.2, 2.2)]
    for (x, y) in uc_main:
        ax.annotate('', xy=(x - 1.5, y), xytext=(1.8, 5.5),
                    arrowprops=dict(arrowstyle='->', color='black', lw=1.2), zorder=6)

    # <<include>> connections
    includes = [
        ((6.2, 7.5), (10.5, 7.5), '<<include>>'),
        ((6.2, 7.5), (10.5, 6.0), '<<include>>'),
        ((6.2, 6.2), (10.5, 6.0), '<<include>>'),
        ((6.2, 6.2), (10.5, 7.5), '<<include>>'),
        ((6.2, 4.9), (10.5, 4.5), '<<include>>'),
        ((6.2, 4.9), (10.5, 8.25), None),
        ((6.2, 3.6), (10.5, 3.0), '<<include>>'),
        ((6.2, 3.6), (10.5, 6.0), '<<include>>'),
        ((6.2, 2.2), (10.5, 6.0), '<<include>>'),
    ]

    for (x1, y1), (x2, y2), lbl in includes:
        ax.annotate('', xy=(x2 - 1.5, y2), xytext=(x1 + 1.5, y1),
                    arrowprops=dict( arrowstyle='->', color='gray', lw=1.0), zorder=2)
        if lbl:
            mx, my = (x1 + x2) / 2, (y1 + y2) / 2
            ax.text(mx, my, lbl, ha='center', va='bottom', fontsize=6.5,
                    color='gray', style='italic')

    buf = _buf()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def make_sequence_login():
    """Sequence diagram: login by password."""
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_facecolor('white')
    fig.patch.set_facecolor('white')

    lifelines = [
        (1.0,  'Пользователь'),
        (3.0,  'LoginFragment'),
        (5.0,  'LoginViewModel'),
        (7.0,  'LoginRepository'),
        (9.2,  'AuthController'),
        (11.4, 'AuthService'),
        (13.2, 'JwtService'),
    ]

    top_y = 7.5
    bot_y = 0.3

    for x, label in lifelines:
        ax.text(x, top_y + 0.2, label, ha='center', va='bottom', fontsize=7.5,
                fontweight='bold', wrap=True, multialignment='center')
        rect = mpatches.FancyBboxPatch((x - 0.45, top_y - 0.15), 0.9, 0.3,
                                        boxstyle="square,pad=0", linewidth=1,
                                        edgecolor='black', facecolor='lightblue', zorder=3)
        ax.add_patch(rect)
        ax.plot([x, x], [top_y - 0.15, bot_y], 'k--', linewidth=0.8, zorder=1)

    msgs = [
        (1.0, 3.0, 7.2, 'login(phone, password)', False),
        (3.0, 5.0, 6.7, 'login(phone, password)', False),
        (5.0, 7.0, 6.2, 'login(phone, password)', False),
        (7.0, 9.2, 5.7, 'POST /api/auth/sign-in', False),
        (9.2, 11.4, 5.2, 'signIn(request)', False),
        (11.4, 13.2, 4.7, 'generateToken(user)\ngenerateRefreshToken(user)', False),
        (13.2, 11.4, 4.1, 'JwtAuthenticationResponse', True),
        (11.4, 9.2,  3.6, 'JwtAuthenticationResponse', True),
        (9.2,  7.0,  3.1, 'AuthTokens', True),
        (7.0,  5.0,  2.6, 'Result.Success', True),
        (5.0,  3.0,  2.1, 'LoginResult(success)', True),
        (3.0,  1.0,  1.5, 'startActivity(MainActivity)', True),
    ]

    activation = {
        3.0:  (6.8, 6.5),
        5.0:  (6.5, 2.3),
        7.0:  (6.0, 2.8),
        9.2:  (5.5, 3.3),
        11.4: (5.0, 3.8),
        13.2: (4.5, 4.3),
    }
    for x, (y_start, y_end) in activation.items():
        rect = mpatches.FancyBboxPatch((x - 0.08, y_end), 0.16, y_start - y_end,
                                        boxstyle="square,pad=0", linewidth=0.8,
                                        edgecolor='black', facecolor='white', zorder=2)
        ax.add_patch(rect)

    for x1, x2, y, label, is_return in msgs:
        style = '<-' if is_return else '->'
        lw    = 0.8
        clr   = 'gray' if is_return else 'black'
        ls    = '--' if is_return else '-'
        ax.annotate('', xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle=style, color=clr, lw=lw, linestyle=ls), zorder=4)
        mx = (x1 + x2) / 2
        ax.text(mx, y + 0.08, label, ha='center', va='bottom', fontsize=6.5,
                color=clr, multialignment='center', style='italic' if is_return else 'normal')

    buf = _buf()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def make_sequence_register():
    """Sequence diagram: registration."""
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_facecolor('white')
    fig.patch.set_facecolor('white')

    lifelines = [
        (0.7,  'Пользователь'),
        (2.5,  'SignUpAuth\nFragment'),
        (4.3,  'SignUpCompletion\nFragment'),
        (6.1,  'LoginViewModel'),
        (8.0,  'LoginRepository'),
        (10.0, 'AuthController'),
        (12.0, 'AuthService /\nUserService'),
        (13.8, 'JwtService'),
    ]
    top_y, bot_y = 7.5, 0.3
    for x, label in lifelines:
        ax.text(x, top_y + 0.2, label, ha='center', va='bottom', fontsize=7,
                fontweight='bold', multialignment='center')
        rect = mpatches.FancyBboxPatch((x - 0.45, top_y - 0.15), 0.9, 0.3,
                                        boxstyle="square,pad=0", linewidth=1,
                                        edgecolor='black', facecolor='lightgreen', zorder=3)
        ax.add_patch(rect)
        ax.plot([x, x], [top_y - 0.15, bot_y], 'k--', linewidth=0.8, zorder=1)

    msgs = [
        (0.7,  2.5, 7.0, 'Ввод телефона и пароля', False),
        (2.5,  4.3, 6.5, 'navigate(phone, password)', False),
        (0.7,  4.3, 6.0, 'Ввод имени и username', False),
        (4.3,  6.1, 5.5, 'signUp(firstName, lastName,\nusername, phone, password)', False),
        (6.1,  8.0, 5.0, 'signUp(...)', False),
        (8.0, 10.0, 4.5, 'POST /api/auth/sign-up', False),
        (10.0, 12.0, 4.0, 'signUp(request)', False),
        (12.0, 13.8, 3.5, 'generateToken()\ngenerateRefreshToken()', False),
        (13.8, 12.0, 3.0, 'JwtAuthResponse', True),
        (12.0, 10.0, 2.5, 'JwtAuthResponse', True),
        (10.0,  8.0, 2.0, 'AuthTokens', True),
        (8.0,   6.1, 1.5, 'Result.Success', True),
        (6.1,   4.3, 1.0, 'LoginResult(success)', True),
        (4.3,   0.7, 0.5, 'startActivity(MainActivity)', True),
    ]

    for x1, x2, y, label, is_return in msgs:
        style = '<-' if is_return else '->'
        clr   = 'gray' if is_return else 'black'
        ls    = '--' if is_return else '-'
        ax.annotate('', xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle=style, color=clr, lw=0.8, linestyle=ls), zorder=4)
        mx = (x1 + x2) / 2
        ax.text(mx, y + 0.08, label, ha='center', va='bottom', fontsize=6.5,
                color=clr, multialignment='center', style='italic' if is_return else 'normal')

    buf = _buf()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def make_sequence_passkey():
    """Sequence diagram: passkey authentication."""
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_facecolor('white')
    fig.patch.set_facecolor('white')

    lifelines = [
        (0.8,  'Пользователь'),
        (2.8,  'WelcomeAuth\nFragment'),
        (4.8,  'Passkey\nAuthManager'),
        (6.8,  'LoginRepository'),
        (8.8,  'Credential\nManager'),
        (10.8, 'PasskeyAuth\nController'),
        (12.8, 'PasskeyService'),
    ]
    top_y, bot_y = 7.5, 0.3
    for x, label in lifelines:
        ax.text(x, top_y + 0.2, label, ha='center', va='bottom', fontsize=7,
                fontweight='bold', multialignment='center')
        rect = mpatches.FancyBboxPatch((x - 0.45, top_y - 0.15), 0.9, 0.3,
                                        boxstyle="square,pad=0", linewidth=1,
                                        edgecolor='black', facecolor='#ffe4b5', zorder=3)
        ax.add_patch(rect)
        ax.plot([x, x], [top_y - 0.15, bot_y], 'k--', linewidth=0.8, zorder=1)

    msgs = [
        (0.8,  2.8, 7.1, 'Нажать «Войти через Passkey»', False),
        (2.8,  4.8, 6.7, 'authenticate(activity, callback)', False),
        (4.8,  6.8, 6.3, 'startPasskeyAuthentication()', False),
        (6.8, 10.8, 5.9, 'POST /passkey/authenticate/start', False),
        (10.8, 12.8, 5.5, 'startAuthentication()', False),
        (12.8, 10.8, 5.1, 'AuthOptions(optionsJson, nonce)', True),
        (10.8,  6.8, 4.7, 'PasskeyAssertionOptions', True),
        (6.8,   4.8, 4.3, 'Result.Success(options)', True),
        (4.8,   8.8, 3.9, 'getCredentialAsync(request)', False),
        (8.8,   0.8, 3.5, 'Биометрическая\nаутентификация', False),
        (0.8,   8.8, 3.1, 'Подтверждение', True),
        (8.8,   4.8, 2.7, 'PublicKeyCredential', True),
        (4.8,   6.8, 2.3, 'finishPasskeyAuthentication(\nnonce, credentialJson)', False),
        (6.8,  10.8, 1.9, 'POST /passkey/authenticate/finish', False),
        (10.8,  4.8, 1.5, 'AuthTokens', True),
        (4.8,   2.8, 1.1, 'onSuccess()', True),
        (2.8,   0.8, 0.6, 'startActivity(MainActivity)', True),
    ]

    for x1, x2, y, label, is_return in msgs:
        style = '<-' if is_return else '->'
        clr   = 'gray' if is_return else 'black'
        ls    = '--' if is_return else '-'
        ax.annotate('', xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle=style, color=clr, lw=0.8, linestyle=ls), zorder=4)
        mx = (x1 + x2) / 2
        ax.text(mx, y + 0.07, label, ha='center', va='bottom', fontsize=6,
                color=clr, multialignment='center', style='italic' if is_return else 'normal')

    buf = _buf()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def make_class_diagram_android():
    """Class diagram – Android auth module."""
    fig, ax = plt.subplots(figsize=(15, 11))
    ax.set_xlim(0, 15)
    ax.set_ylim(0, 11)
    ax.axis('off')
    ax.set_facecolor('white')
    fig.patch.set_facecolor('white')

    def class_box(ax, x, y, w, h, name, attrs, methods, color='#ddeeff'):
        # outer rectangle
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="square,pad=0",
                                        linewidth=1.2, edgecolor='black', facecolor=color, zorder=2)
        ax.add_patch(rect)
        # name section
        name_h = 0.42
        ax.fill_between([x, x+w], [y+h-name_h, y+h-name_h], [y+h, y+h],
                         color='#99bbdd', zorder=3)
        ax.plot([x, x+w], [y+h-name_h, y+h-name_h], 'k-', lw=0.8, zorder=3)
        ax.text(x + w/2, y + h - name_h/2, name, ha='center', va='center',
                fontsize=7.5, fontweight='bold', zorder=4)
        # attributes
        sep_attrs = y + h - name_h - len(attrs) * 0.25
        for i, attr in enumerate(attrs):
            ax.text(x + 0.08, y + h - name_h - 0.22 - i * 0.25, attr,
                    ha='left', va='center', fontsize=6.5, zorder=4)
        ax.plot([x, x+w], [sep_attrs, sep_attrs], 'k-', lw=0.5, zorder=3)
        # methods
        for i, mth in enumerate(methods):
            ax.text(x + 0.08, sep_attrs - 0.22 - i * 0.25, mth,
                    ha='left', va='center', fontsize=6.5, zorder=4)

    def arrow(ax, x1, y1, x2, y2, style='->', color='black', dashed=False):
        ls = '--' if dashed else '-'
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle=style, color=color, lw=1.0, linestyle=ls), zorder=5)

    def label_arrow(ax, x1, y1, x2, y2, label, dashed=False):
        arrow(ax, x1, y1, x2, y2, dashed=dashed)
        ax.text((x1+x2)/2, (y1+y2)/2 + 0.1, label,
                ha='center', va='bottom', fontsize=6, color='gray', style='italic')

    # LoginActivity
    class_box(ax, 0.2, 9.0, 3.0, 1.8, 'LoginActivity',
              ['savedInstanceState: Bundle'],
              ['onCreate(): void', 'hasActiveSession(): bool', 'openMainScreen(): void'])

    # LoginViewModel
    class_box(ax, 4.0, 9.0, 3.2, 1.8, 'LoginViewModel',
              ['loginFormState: LiveData', 'loginResult: LiveData', 'signUpResult: LiveData'],
              ['login(phone, password): void', 'signUp(...): void', 'loginDataChanged(): void'])

    # LoginRepository
    class_box(ax, 8.0, 9.0, 3.5, 1.8, 'LoginRepository',
              ['authApi: AuthApi', 'tokenManager: TokenManager'],
              ['login(): Result', 'signUp(): Result', 'startPasskeyAuth(): Result',
               'finishPasskeyAuth(): Result'])

    # TokenManager
    class_box(ax, 11.8, 9.0, 3.0, 1.8, 'TokenManager',
              ['prefs: SharedPreferences'],
              ['saveTokens(): void', 'getAccessToken(): String', 'getRefreshToken(): String',
               'isAccessTokenExpired(): bool', 'clearTokens(): void'])

    # Fragments group
    class_box(ax, 0.2, 6.8, 2.2, 1.9, 'WelcomeAuth\nFragment',
              ['passkeyAuthManager'],
              ['onViewCreated()', 'startPasskeyFlow()'])

    class_box(ax, 2.6, 6.8, 2.2, 1.9, 'LoginFragment',
              ['loginViewModel', 'binding'],
              ['onViewCreated()', 'updateUiWithUser()', 'showLoginFailed()'])

    class_box(ax, 5.0, 6.8, 2.2, 1.9, 'SignUpAuth\nFragment',
              ['binding'],
              ['onViewCreated()'])

    class_box(ax, 7.4, 6.8, 2.4, 1.9, 'SignUpCompletion\nInternal',
              ['loginViewModel', 'binding'],
              ['onViewCreated()'])

    class_box(ax, 10.0, 6.8, 2.6, 1.9, 'SignUpCompletion\n3rdParty',
              ['passkeyAuthManager', 'binding'],
              ['onViewCreated()', 'startPasskeyReg()'])

    # PasskeyAuthManager
    class_box(ax, 0.2, 4.3, 3.2, 2.2, 'PasskeyAuthManager',
              ['loginRepository', 'credentialManager', 'ioExecutor', 'mainHandler'],
              ['authenticate(): void', 'register(): void', 'finishAuthentication(): void',
               'finishRegistration(): void'])

    # AuthInterceptor
    class_box(ax, 4.2, 4.3, 3.0, 2.2, 'AuthInterceptor',
              ['tokenManager: TokenManager', 'authApi: AuthApi'],
              ['intercept(): Response', 'getValidToken(): String', 'refreshTokens(): String'])

    # AuthApi
    class_box(ax, 8.0, 4.3, 3.0, 2.2, '«interface»\nAuthApi',
              [],
              ['login(): Call<AuthTokens>', 'signUp(): Call<AuthTokens>',
               'refreshToken(): Call<AuthTokens>', 'getPasskeyOptions(): Call',
               'completePasskey(): Call', 'startPasskeyAuth(): Call',
               'finishPasskeyAuth(): Call'])

    # ApiClient
    class_box(ax, 11.5, 4.3, 3.0, 1.5, 'ApiClient',
              ['retrofit: Retrofit'],
              ['getInstance(): Retrofit', 'api(): ApiService'])

    # LoggedInUserView
    class_box(ax, 0.2, 2.2, 2.5, 1.3, 'LoggedInUserView',
              ['displayName: String'],
              ['getDisplayName(): String'])

    # LoginFormState
    class_box(ax, 3.0, 2.2, 2.5, 1.3, 'LoginFormState',
              ['phoneError: Integer', 'passwordError: Integer', 'isDataValid: bool'],
              ['isDataValid(): bool', 'getPhoneError(): Integer'])

    # LoginResult
    class_box(ax, 6.0, 2.2, 2.5, 1.3, 'LoginResult',
              ['success: LoggedInUserView', 'error: Integer'],
              ['getSuccess()', 'getError()'])

    # AuthRepositoryFactory
    class_box(ax, 9.0, 2.2, 2.8, 1.3, 'AuthRepositoryFactory',
              [],
              ['create(context): LoginRepository'])

    # LoginViewModelFactory
    class_box(ax, 12.0, 2.2, 2.8, 1.3, 'LoginViewModelFactory',
              ['context: Context'],
              ['create(modelClass): T'])

    # Relationships
    arrow(ax, 3.2, 9.9, 4.0, 9.9)           # LoginActivity -> LoginViewModel (via Factory)
    arrow(ax, 7.2, 9.9, 8.0, 9.9)           # LoginViewModel -> LoginRepository
    arrow(ax, 11.5, 9.9, 11.8, 9.9)         # LoginRepository -> TokenManager
    arrow(ax, 9.75, 9.0, 9.75, 6.5, dashed=True)  # LoginRepository (usage)
    arrow(ax, 2.6, 7.75, 4.0, 9.9, dashed=True, color='gray')   # LoginFragment -> LoginViewModel
    arrow(ax, 7.4, 7.75, 5.2, 9.9, dashed=True, color='gray')   # SignUpComp -> LoginViewModel
    arrow(ax, 1.3, 6.8, 1.8, 6.5)           # WelcomeAuthFr -> PasskeyAuthManager
    arrow(ax, 10.0+1.3, 6.8, 1.8, 5.5, dashed=True, color='gray')
    arrow(ax, 1.8, 4.3, 8.5, 6.5, dashed=True, color='gray')  # PasskeyAuthMgr -> LoginRepo
    arrow(ax, 5.7, 4.3, 8.0, 5.5, dashed=True, color='gray')  # AuthInterceptor -> AuthApi
    arrow(ax, 4.5, 4.3, 4.5, 3.5, dashed=True, color='gray')  # AuthInterceptor -> TokenManager
    arrow(ax, 12.0, 2.85, 9.8, 2.85)        # LoginVMFactory -> LoginRepo (creates)
    arrow(ax, 9.0, 2.85, 6.5, 2.55, dashed=True, color='gray')  # AuthRepoFactory -> LoginRepo

    buf = _buf()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def make_class_diagram_backend():
    """Class diagram – Backend auth module."""
    fig, ax = plt.subplots(figsize=(15, 9))
    ax.set_xlim(0, 15)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.set_facecolor('white')
    fig.patch.set_facecolor('white')

    def class_box(ax, x, y, w, h, name, attrs, methods, color='#ddeedd'):
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="square,pad=0",
                                        linewidth=1.2, edgecolor='black', facecolor=color, zorder=2)
        ax.add_patch(rect)
        name_h = 0.42
        ax.fill_between([x, x+w], [y+h-name_h, y+h-name_h], [y+h, y+h],
                         color='#88bb99', zorder=3)
        ax.plot([x, x+w], [y+h-name_h, y+h-name_h], 'k-', lw=0.8, zorder=3)
        ax.text(x + w/2, y + h - name_h/2, name, ha='center', va='center',
                fontsize=7.5, fontweight='bold', zorder=4)
        sep_attrs = y + h - name_h - len(attrs) * 0.25
        for i, attr in enumerate(attrs):
            ax.text(x + 0.08, y + h - name_h - 0.22 - i * 0.25, attr,
                    ha='left', va='center', fontsize=6.5, zorder=4)
        ax.plot([x, x+w], [sep_attrs, sep_attrs], 'k-', lw=0.5, zorder=3)
        for i, mth in enumerate(methods):
            ax.text(x + 0.08, sep_attrs - 0.22 - i * 0.25, mth,
                    ha='left', va='center', fontsize=6.5, zorder=4)

    def arrow(ax, x1, y1, x2, y2, dashed=False):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='black', lw=1.0,
                                   linestyle='--' if dashed else '-'), zorder=5)

    # AuthController
    class_box(ax, 0.2, 6.8, 3.5, 1.8, 'AuthController',
              ['authenticationService: AuthenticationService'],
              ['signUp(request): ResponseEntity', 'signIn(request): ResponseEntity',
               'refresh(request): ResponseEntity'])

    # AuthenticationService
    class_box(ax, 4.5, 6.8, 3.6, 1.8, 'AuthenticationService',
              ['userService', 'jwtService', 'passwordEncoder', 'authManager'],
              ['signUp(request): JwtResponse', 'signIn(request): JwtResponse',
               'refresh(request): JwtResponse', 'buildResponse(user): JwtResponse'])

    # JwtService
    class_box(ax, 9.0, 6.8, 5.5, 1.8, 'JwtService',
              ['jwtSecret: String', 'expirationMs: long', 'refreshExpirationMs: long'],
              ['generateToken(userDetails): String', 'generateRefreshToken(): String',
               'isTokenValid(): bool', 'extractSubject(): String',
               'generateRegistrationToken(): String', 'parseRegistrationToken(): Claims'])

    # JwtAuthenticationFilter
    class_box(ax, 0.2, 4.5, 3.5, 2.0, 'JwtAuthenticationFilter',
              ['jwtService: JwtService', 'userService: UserService'],
              ['doFilterInternal(req, res, chain): void'])

    # SecurityConfig
    class_box(ax, 4.5, 4.5, 3.5, 2.0, 'SecurityConfig',
              ['jwtAuthFilter: JwtAuthFilter', 'userService: UserService'],
              ['securityFilterChain(): SecurityFilterChain',
               'passwordEncoder(): PasswordEncoder',
               'authenticationProvider(): AuthProvider',
               'authenticationManager(): AuthManager'])

    # PasskeyAuthController
    class_box(ax, 9.0, 4.5, 5.5, 2.0, 'PasskeyAuthController',
              ['passkeyService: PasskeyService', 'jwtService: JwtService'],
              ['options(request): ResponseEntity', 'complete(nonce, cred): ResponseEntity',
               'authenticateStart(): ResponseEntity', 'authenticateFinish(): ResponseEntity'])

    # User (entity)
    class_box(ax, 0.2, 2.0, 3.5, 2.2, 'User (Entity)',
              ['id: UUID', 'firstName: String', 'lastName: String',
               'username: String', 'phone: String', 'password: String', 'avatarUrl: String'],
              ['getUsername(): String (phone)', 'getHandle(): String (username)',
               'getAuthorities(): Collection'])

    # UserService
    class_box(ax, 4.5, 2.0, 3.5, 2.2, 'UserService',
              ['repository: UserRepository'],
              ['registerUser(): User', 'userDetailsService(): UDS',
               'findByEmail(): Optional<User>', 'findById(): Optional<User>'])

    # UserRepository
    class_box(ax, 9.0, 2.0, 4.5, 2.2, 'UserRepository (JPA)',
              [],
              ['findByPhone(): Optional<User>', 'findByEmail(): Optional<User>',
               'existsByPhone(): boolean', 'existsByUsername(): boolean',
               'existsByEmail(): boolean'])

    # DTOs
    class_box(ax, 0.2, 0.2, 2.5, 1.5, 'SignUpRequest',
              ['firstName', 'lastName', 'username', 'phone', 'password'], [])

    class_box(ax, 3.0, 0.2, 2.5, 1.5, 'SignInRequest',
              ['phone: String', 'password: String'], [])

    class_box(ax, 5.7, 0.2, 3.2, 1.5, 'JwtAuthenticationResponse',
              ['accessToken: String', 'refreshToken: String'], [])

    class_box(ax, 9.2, 0.2, 2.5, 1.5, 'RefreshTokenRequest',
              ['refreshToken: String'], [])

    class_box(ax, 12.0, 0.2, 2.8, 1.5, 'PasskeyCredential',
              ['credentialId: String', 'userId: UUID', 'publicKey: byte[]'], [])

    # Arrows
    arrow(ax, 3.7, 7.7, 4.5, 7.7)             # AuthController -> AuthService
    arrow(ax, 8.1, 7.7, 9.0, 7.7)             # AuthService -> JwtService
    arrow(ax, 8.1, 7.3, 8.1, 3.1, dashed=True)
    arrow(ax, 5.2, 4.5, 5.5, 6.8, dashed=True)  # SecurityConfig -> AuthFilter
    arrow(ax, 0.2+1.7, 4.5, 0.2+1.7, 6.8, dashed=True)  # Filter -> JwtService
    arrow(ax, 4.5+1.7, 2.0, 4.5+1.7, 4.5, dashed=True)  # AuthService -> UserService
    arrow(ax, 4.5+1.7, 2.0, 8.1, 7.3, dashed=True)
    arrow(ax, 8.1, 2.0, 9.0, 2.0, dashed=True)  # UserService -> UserRepository
    arrow(ax, 9.0+2.2, 4.5, 9.0+2.2, 6.8, dashed=True)  # PasskeyCtrl -> JwtService

    buf = _buf()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def set_doc_styles(doc):
    sec = doc.sections[0]
    sec.left_margin   = LEFT_MARGIN
    sec.right_margin  = RIGHT_MARGIN
    sec.top_margin    = TOP_MARGIN
    sec.bottom_margin = BOT_MARGIN

    normal = doc.styles['Normal']
    normal.font.name  = FONT_NAME
    normal.font.size  = FONT_SIZE
    from docx.oxml import OxmlElement
    rPr = normal.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    FONT_NAME)
    rFonts.set(qn('w:hAnsi'),    FONT_NAME)
    rFonts.set(qn('w:cs'),       FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rPr.insert(0, rFonts)


def add_title_page(doc):
    def center(text, bold=False, size=None):
        p = doc.add_paragraph()
        r = p.add_run(text)
        _set_run_font(r, bold=bold, size=size or FONT_SIZE)
        p.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before      = Pt(0)
        p.paragraph_format.space_after       = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        return p

    center("Министерство науки и высшего образования Российской Федерации")
    center("Федеральное государственное бюджетное образовательное учреждение")
    center("высшего образования")
    center("ИРКУТСКИЙ НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ", bold=True)
    center("Институт информационных технологий и анализа данных")
    doc.add_paragraph()
    doc.add_paragraph()
    center("Допускаю к защите")
    center("Руководитель ____________ _________________")
    doc.add_paragraph()
    doc.add_paragraph()
    center("Разработка модуля аутентификации и авторизации\nмобильного мессенджера Flux Messenger",
           bold=True, size=Pt(16))
    doc.add_paragraph()
    center("ПОЯСНИТЕЛЬНАЯ ЗАПИСКА")
    center("к курсовому проекту по дисциплине")
    center("Объектно-ориентированное программирование")
    doc.add_paragraph()
    center("1.023.00.00 ПЗ")
    doc.add_paragraph()
    doc.add_paragraph()
    center("Выполнил студент группы ________")
    center("______________________________")
    doc.add_paragraph()
    center("Нормоконтроль ______________________")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    center("Иркутск 2025 г.")
    page_break(doc)


def add_toc(doc):
    add_heading(doc, "Содержание", level=1)
    entries = [
        ("Введение", ""),
        ("1  Анализ предметной области", ""),
        ("1.1  Описание предметной области", ""),
        ("1.2  Диаграмма AS IS", ""),
        ("1.3  Постановка задачи", ""),
        ("2  Проектирование", ""),
        ("2.1  Диаграмма TO BE", ""),
        ("2.3  UML диаграмма use case", ""),
        ("2.4  Проектирование графического интерфейса", ""),
        ("3  Реализация", ""),
        ("3.1  Использование технологий и библиотек", ""),
        ("3.2  Диаграммы последовательности", ""),
        ("3.3  Спецификация", ""),
        ("3.4  UML диаграмма классов", ""),
        ("3.5  Взаимодействие с БД", ""),
        ("3.6  Тестирование", ""),
        ("3.7  Руководство пользователя", ""),
        ("Заключение", ""),
        ("Список литературы", ""),
    ]
    for title, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before      = Pt(0)
        p.paragraph_format.space_after       = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        tab_stops = p.paragraph_format.tab_stops
        from docx.shared import Cm as _Cm
        tab_stops.add_tab_stop(_Cm(15.5))
        run = p.add_run(title)
        _set_run_font(run)
    page_break(doc)


def build_report():
    doc = Document()
    set_doc_styles(doc)
    add_title_page(doc)
    add_toc(doc)

    # ── ВВЕДЕНИЕ ─────────────────────────────────────────────────────────────
    add_heading(doc, "Введение", level=1)

    add_para(doc,
        "В условиях повсеместного распространения мобильных технологий мессенджеры "
        "превратились в один из ключевых инструментов межличностной коммуникации. "
        "Безопасность и удобство входа в приложение непосредственно определяют "
        "доверие пользователей и уровень защиты персональных данных. Традиционная "
        "парольная аутентификация обладает рядом известных уязвимостей: подверженность "
        "фишинговым атакам, перебору паролей и компрометации баз данных. Переход к "
        "беспарольным технологиям — в частности, к стандарту WebAuthn/Passkey — "
        "позволяет существенно повысить уровень безопасности без снижения удобства "
        "пользователя. Актуальность разработки системы аутентификации для мобильного "
        "мессенджера обусловлена необходимостью обеспечить надёжную защиту переписки "
        "пользователей при использовании как классического, так и биометрического "
        "метода входа.")

    add_para(doc,
        "Целью курсовой работы является разработка модуля аутентификации и авторизации "
        "для мобильного Android-мессенджера Flux Messenger с поддержкой JWT-токенов "
        "и технологии Passkey (WebAuthn).")

    add_para(doc, "Для достижения поставленной цели были сформулированы следующие задачи:")

    tasks = [
        "изучить и проанализировать методы аутентификации в современных мобильных приложениях;",
        "спроектировать архитектуру модуля аутентификации с учётом паттерна MVVM;",
        "разработать серверную часть аутентификации на основе Spring Boot, Spring Security и JWT;",
        "реализовать поддержку беспарольного входа посредством WebAuthn/Passkey;",
        "разработать клиентский модуль Android-приложения с безопасным хранением токенов;",
        "провести тестирование ключевых компонентов модуля."
    ]
    for t in tasks:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(t)
        _set_run_font(run)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    add_para(doc,
        "Объектом исследования является система аутентификации и авторизации пользователей "
        "мобильного мессенджера Flux Messenger.")
    add_para(doc,
        "Предметом исследования служат методы и технологии реализации многофакторной "
        "аутентификации в Android-приложениях с серверной частью на основе Spring Boot.")

    page_break(doc)

    # ── РАЗДЕЛ 1 ─────────────────────────────────────────────────────────────
    add_heading(doc, "Анализ предметной области", level=1, numbered="1")

    # 1.1
    add_heading(doc, "Описание предметной области", level=2, numbered="1.1")

    add_para(doc,
        "Flux Messenger — мобильное Android-приложение, предназначенное для обмена "
        "текстовыми сообщениями и медиафайлами в режиме реального времени. Ключевым "
        "компонентом любого мессенджера является подсистема аутентификации, обеспечивающая "
        "идентификацию пользователей и защиту их переписки от несанкционированного доступа.")

    add_para(doc,
        "Предметная область охватывает взаимодействие пользователей с системой верификации "
        "личности, хранение и валидацию криптографических токенов, а также управление "
        "сессиями. Основные понятия предметной области:")

    terms = [
        ("Аутентификация (Authentication)",
         "процесс проверки подлинности пользователя посредством предоставленных учётных данных."),
        ("Авторизация (Authorization)",
         "предоставление аутентифицированному пользователю прав доступа к ресурсам приложения."),
        ("JWT (JSON Web Token)",
         "открытый стандарт (RFC 7519) для компактного и самодостаточного представления "
         "утверждений между сторонами в виде JSON-объекта, подписанного криптографическим ключом."),
        ("Access Token (токен доступа)",
         "краткосрочный JWT (время жизни 15 минут), прикладываемый к каждому запросу к API "
         "для подтверждения аутентичности пользователя."),
        ("Refresh Token (токен обновления)",
         "долгосрочный JWT (время жизни 7 дней), используемый для получения нового токена "
         "доступа по истечении срока его действия без повторного ввода пароля."),
        ("Passkey (WebAuthn)",
         "стандарт аутентификации W3C, основанный на криптографических ключевых парах: "
         "приватный ключ хранится на устройстве под защитой биометрии или PIN-кода, "
         "публичный ключ регистрируется на сервере."),
        ("EncryptedSharedPreferences",
         "компонент AndroidX Security, обеспечивающий хранение данных в зашифрованном "
         "виде с использованием алгоритмов AES-256-SIV и AES-256-GCM."),
    ]
    for term, defn in terms:
        p = doc.add_paragraph()
        run_bold = p.add_run(term + " — ")
        _set_run_font(run_bold, bold=True)
        run_text = p.add_run(defn)
        _set_run_font(run_text)
        _fmt_para(p)

    add_para(doc,
        "Основными субъектами предметной области являются: незарегистрированный пользователь, "
        "имеющий доступ только к экранам входа и регистрации; аутентифицированный пользователь, "
        "получающий доступ к функциям мессенджера; сервер аутентификации, проверяющий "
        "учётные данные и генерирующий токены.")

    # 1.2
    add_heading(doc, "Диаграмма AS IS", level=2, numbered="1.2")

    add_para(doc,
        "На рисунке 1.2 представлена диаграмма AS IS, описывающая текущий процесс "
        "аутентификации пользователей в мессенджерах до внедрения разрабатываемого решения.")

    p = doc.add_paragraph()
    run = p.add_run(
        "[МЕСТО ДЛЯ ДИАГРАММЫ AS IS — вставьте диаграмму здесь]"
    )
    _set_run_font(run)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    add_figure_caption(doc, "1.2", "Диаграмма AS IS — текущий процесс аутентификации")

    add_para(doc,
        "Традиционный процесс аутентификации в мессенджерах предполагает ручной ввод "
        "логина и пароля при каждом новом сеансе. Пользователь несёт ответственность "
        "за надёжность пароля и его безопасное хранение, что на практике нередко приводит "
        "к использованию слабых или повторяющихся паролей. Отсутствие автоматического "
        "обновления сессий вынуждает пользователя повторно проходить аутентификацию, "
        "снижая удобство использования приложения.")

    # 1.3
    add_heading(doc, "Постановка задачи", level=2, numbered="1.3")

    add_para(doc,
        "Необходимо разработать модуль аутентификации и авторизации для мобильного "
        "мессенджера Flux Messenger. Модуль должен состоять из двух частей: серверной "
        "(бэкенд) и клиентской (Android-приложение).")

    add_para(doc, "Серверная часть должна реализовывать следующий функционал:")

    server_funcs = [
        "POST /api/auth/sign-up — регистрация нового пользователя по номеру телефона, "
        "имени, имени пользователя и паролю;",
        "POST /api/auth/sign-in — вход существующего пользователя по номеру телефона и паролю;",
        "POST /api/auth/refresh — обновление истёкшего токена доступа по токену обновления;",
        "POST /api/auth/passkey/options — формирование параметров для регистрации Passkey;",
        "POST /api/auth/passkey/complete — завершение регистрации Passkey и выдача токенов;",
        "POST /api/auth/passkey/authenticate/start — инициализация аутентификации по Passkey;",
        "POST /api/auth/passkey/authenticate/finish — верификация ответа аутентификатора "
        "и выдача токенов.",
    ]
    for f in server_funcs:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f)
        _set_run_font(run)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    add_para(doc, "Клиентская часть должна реализовывать следующий функционал:")

    client_funcs = [
        "экран приветствия с выбором метода аутентификации (пароль или Passkey);",
        "форму входа по номеру телефона и паролю с валидацией вводимых данных;",
        "двухэтапную регистрацию: ввод телефона и пароля, затем личных данных;",
        "поддержку аутентификации и регистрации через Passkey (WebAuthn);",
        "безопасное хранение токенов в EncryptedSharedPreferences (AES-256-GCM);",
        "автоматическое добавление Bearer-токена к HTTP-запросам;",
        "автоматическое обновление истёкшего токена доступа посредством OkHttp-перехватчика.",
    ]
    for f in client_funcs:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f)
        _set_run_font(run)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    page_break(doc)

    # ── РАЗДЕЛ 2 ─────────────────────────────────────────────────────────────
    add_heading(doc, "Проектирование", level=1, numbered="2")

    # 2.1
    add_heading(doc, "Диаграмма TO BE", level=2, numbered="2.1")

    add_para(doc,
        "На рисунке 2.1 представлена диаграмма TO BE, отражающая спроектированный процесс "
        "аутентификации в мессенджере Flux Messenger после внедрения разрабатываемого модуля.")

    p = doc.add_paragraph()
    run = p.add_run(
        "[МЕСТО ДЛЯ ДИАГРАММЫ TO BE — вставьте диаграмму здесь]"
    )
    _set_run_font(run)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    add_figure_caption(doc, "2.1", "Диаграмма TO BE — спроектированный процесс аутентификации")

    add_para(doc,
        "Разработанный процесс аутентификации включает следующие усовершенствования: "
        "автоматическое обновление токена доступа без участия пользователя, поддержку "
        "биометрической аутентификации через Passkey, а также безопасное зашифрованное "
        "хранение токенов на устройстве. Пользователь однократно проходит регистрацию "
        "или вход, после чего приложение самостоятельно управляет жизненным циклом сессии.")

    # 2.3
    add_heading(doc, "UML диаграмма use case", level=2, numbered="2.3")

    add_para(doc,
        "В соответствии с рисунком 2.2 представлена UML-диаграмма вариантов использования "
        "модуля аутентификации. Актором выступает пользователь — субъект, взаимодействующий "
        "с системой через интерфейс мобильного приложения.")

    uc_img = make_usecase_diagram()
    embed_image(doc, uc_img, width_cm=15.0)
    add_figure_caption(doc, "2.2", "UML-диаграмма вариантов использования модуля аутентификации")

    add_para(doc,
        "Диаграмма включает следующие варианты использования:")
    uc_items = [
        ("Войти по паролю", "пользователь вводит номер телефона и пароль; система валидирует "
         "данные и при успехе выдаёт пару токенов, которые сохраняются на устройстве."),
        ("Зарегистрироваться по паролю", "пользователь заполняет двухэтапную форму: ввод "
         "телефона и пароля, затем — личных данных (имя, фамилия, имя пользователя); "
         "сервер регистрирует пользователя и выдаёт токены."),
        ("Войти с Passkey", "система запрашивает на сервере параметры аутентификации, "
         "передаёт их в CredentialManager, пользователь подтверждает вход биометрией; "
         "устройство возвращает криптографический ответ, сервер верифицирует его."),
        ("Зарегистрировать Passkey", "пользователь вводит номер телефона; сервер формирует "
         "параметры для создания ключа; пользователь подтверждает биометрией; "
         "CredentialManager генерирует ключевую пару; публичный ключ регистрируется на сервере."),
        ("Перейти к главному экрану", "общий результат успешной аутентификации или "
         "регистрации — переход в MainActivity с сохранёнными токенами."),
    ]
    for name, desc in uc_items:
        p = doc.add_paragraph()
        run_b = p.add_run(f"«{name}» — ")
        _set_run_font(run_b, bold=True)
        run_d = p.add_run(desc)
        _set_run_font(run_d)
        _fmt_para(p)

    # 2.4
    add_heading(doc, "Проектирование графического интерфейса", level=2, numbered="2.4")

    add_para(doc,
        "Навигация между экранами аутентификации реализована с помощью компонента "
        "AndroidX Navigation и описана в графе nav_auth.xml. Начальным экраном является "
        "WelcomeAuthFragment. В таблице 2.1 представлено описание всех экранов "
        "модуля аутентификации.")

    add_table_caption(doc, "2.1", "Экраны модуля аутентификации")
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    headers = ["Экран (фрагмент)", "Назначение", "Элементы управления"]
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            _set_run_font(run, bold=True, size=Pt(12))

    screens = [
        ("WelcomeAuthFragment",
         "Стартовый экран выбора метода входа",
         "Кнопка «Войти», кнопка «Зарегистрироваться», кнопка «Войти через Passkey»"),
        ("LoginFragment",
         "Форма входа по номеру телефона и паролю",
         "Поле ввода телефона (PhoneInputView), поле пароля (PasswordInputView), "
         "кнопка «Войти», ссылка «Зарегистрироваться»"),
        ("SignUpAuthFragment",
         "Первый этап регистрации: данные для входа",
         "Поле телефона, поле пароля, поле подтверждения пароля, кнопка «Далее»"),
        ("SignUpCompletionInternalFragment",
         "Второй этап регистрации: личные данные",
         "Поле имени (AvatarInputView с полями First/Last Name), поле username, "
         "кнопка «Готово»"),
        ("SignUpCompletion3rdPartyFragment",
         "Регистрация через Passkey: личные данные и телефон",
         "AvatarInputView, поле username, поле телефона, "
         "кнопка выбора фото аватара, кнопка «Войти»"),
    ]
    for row_data in screens:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            for run in row[i].paragraphs[0].runs:
                _set_run_font(run, size=Pt(11))

    doc.add_paragraph()
    add_para(doc,
        "Граф навигации описывает следующие переходы: из WelcomeAuthFragment "
        "пользователь может перейти к LoginFragment (кнопка «Войти»), к "
        "SignUpAuthFragment (кнопка «Зарегистрироваться») или к "
        "SignUpCompletion3rdPartyFragment (кнопка «Passkey»). Из LoginFragment "
        "доступен переход к SignUpAuthFragment. Из SignUpAuthFragment — к "
        "SignUpCompletionInternalFragment (кнопка «Далее») или обратно к LoginFragment.")

    page_break(doc)

    # ── РАЗДЕЛ 3 ─────────────────────────────────────────────────────────────
    add_heading(doc, "Реализация", level=1, numbered="3")

    # 3.1
    add_heading(doc, "Использование технологий и библиотек", level=2, numbered="3.1")

    add_para(doc,
        "В таблице 3.1 представлены технологии и библиотеки, использованные при разработке "
        "серверной и клиентской частей модуля аутентификации.")

    add_table_caption(doc, "3.1", "Технологии и библиотеки модуля аутентификации")
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    for i, h in enumerate(["Технология / библиотека", "Версия", "Назначение"]):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            _set_run_font(run, bold=True, size=Pt(11))

    tech_rows = [
        ("Spring Boot", "4.x", "Основной фреймворк серверной части"),
        ("Spring Security", "6.x", "Аутентификация и авторизация на стороне сервера"),
        ("JJWT (Java JWT)", "0.11+", "Генерация и валидация JWT-токенов"),
        ("Spring Security WebAuthn", "6.x", "Поддержка протокола WebAuthn/Passkey"),
        ("Spring Data JPA + Hibernate", "3.x", "ORM для взаимодействия с БД PostgreSQL"),
        ("Lombok", "1.18+", "Генерация шаблонного Java-кода"),
        ("AndroidX Credentials API", "1.3+", "Работа с Passkey/WebAuthn на Android"),
        ("Retrofit 3 + Gson", "3.x", "HTTP-клиент и JSON-сериализация (Android)"),
        ("OkHttp3", "4.x", "HTTP-перехватчики, повторные запросы (Android)"),
        ("EncryptedSharedPreferences", "1.1+", "Зашифрованное хранилище токенов (AES-256-GCM)"),
        ("AndroidX Navigation", "2.x", "Навигация между фрагментами (Android)"),
        ("Glide", "4.x", "Загрузка и отображение изображений (Android)"),
        ("AndroidX ViewModel + LiveData", "2.x", "Архитектурные компоненты MVVM (Android)"),
    ]
    for row_data in tech_rows:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            for run in row[i].paragraphs[0].runs:
                _set_run_font(run, size=Pt(11))

    doc.add_paragraph()
    add_para(doc,
        "На серверной стороне аутентификация реализована следующим образом: Spring Security "
        "перехватывает входящие запросы через JwtAuthenticationFilter, извлекает Bearer-токен "
        "из заголовка Authorization, валидирует его через JwtService и устанавливает "
        "контекст безопасности. Конечные точки /api/auth/** открыты для всех запросов, "
        "остальные требуют действующего токена.")

    add_para(doc,
        "На клиентской стороне применяется архитектурный паттерн MVVM: LoginViewModel "
        "выступает посредником между UI-фрагментами и LoginRepository, который выполняет "
        "HTTP-запросы к серверу через интерфейс AuthApi (Retrofit). Полученные токены "
        "сохраняются в EncryptedSharedPreferences посредством TokenManager. "
        "AuthInterceptor автоматически добавляет токен к каждому запросу и при "
        "получении ответа 401 инициирует обновление токена.")

    # 3.2
    add_heading(doc, "Диаграммы последовательности", level=2, numbered="3.2")

    add_para(doc,
        "Ниже представлены диаграммы последовательности для трёх основных сценариев "
        "аутентификации: вход по паролю, регистрация и вход через Passkey.")

    add_para(doc,
        "В соответствии с рисунком 3.1 показана последовательность взаимодействия "
        "компонентов при входе пользователя по номеру телефона и паролю.")

    seq_login = make_sequence_login()
    embed_image(doc, seq_login, width_cm=15.5)
    add_figure_caption(doc, "3.1", "Диаграмма последовательности: вход по паролю")

    add_para(doc,
        "При нажатии кнопки «Войти» LoginFragment передаёт введённые данные в LoginViewModel, "
        "который делегирует запрос LoginRepository. Репозиторий выполняет синхронный HTTP-запрос "
        "POST /api/auth/sign-in через AuthApi. На сервере AuthController вызывает "
        "AuthenticationService, который через Spring AuthenticationManager верифицирует "
        "учётные данные и посредством JwtService генерирует пару токенов (access + refresh). "
        "Токены возвращаются клиенту, сохраняются в TokenManager, и пользователь "
        "перенаправляется в MainActivity.")

    add_para(doc,
        "В соответствии с рисунком 3.2 показана последовательность взаимодействия "
        "компонентов при регистрации нового пользователя.")

    seq_reg = make_sequence_register()
    embed_image(doc, seq_reg, width_cm=15.5)
    add_figure_caption(doc, "3.2", "Диаграмма последовательности: регистрация пользователя")

    add_para(doc,
        "Регистрация состоит из двух этапов. На первом этапе SignUpAuthFragment получает "
        "номер телефона и пароль с подтверждением, валидирует их и передаёт аргументами "
        "в SignUpCompletionInternalFragment. На втором этапе пользователь вводит имя, "
        "фамилию и имя пользователя. После нажатия «Готово» LoginViewModel.signUp() "
        "отправляет POST /api/auth/sign-up. Сервер через UserService создаёт запись в БД, "
        "а JwtService генерирует токены.")

    add_para(doc,
        "В соответствии с рисунком 3.3 показана последовательность взаимодействия "
        "компонентов при аутентификации через Passkey.")

    seq_pk = make_sequence_passkey()
    embed_image(doc, seq_pk, width_cm=15.5)
    add_figure_caption(doc, "3.3", "Диаграмма последовательности: аутентификация через Passkey")

    add_para(doc,
        "При нажатии кнопки «Войти через Passkey» WelcomeAuthFragment вызывает "
        "PasskeyAuthManager.authenticate(). Менеджер запрашивает у сервера параметры "
        "аутентификации (JSON с challenge) и передаёт их в AndroidX CredentialManager. "
        "CredentialManager инициирует системный диалог биометрической аутентификации. "
        "После подтверждения устройство формирует криптографический ответ, который "
        "LoginRepository отправляет серверу. PasskeyService верифицирует подпись и "
        "возвращает пару токенов. Если пользователь ещё не зарегистрирован "
        "(NoCredentialException), WelcomeAuthFragment перенаправляет к "
        "SignUpCompletion3rdPartyFragment для регистрации Passkey.")

    # 3.3
    add_heading(doc, "Спецификация", level=2, numbered="3.3")

    add_para(doc,
        "В данном разделе приведены спецификации классов, реализованных в рамках "
        "модуля аутентификации. Рассмотрены классы серверной и клиентской частей.")

    add_para(doc, "Серверная часть.", bold=True, indent=True)

    # AuthController spec
    add_para(doc, "Класс AuthController", bold=True, indent=True, space_before=4)
    add_para(doc,
        "REST-контроллер, обрабатывающий запросы аутентификации и регистрации. "
        "Аннотирован @RestController, @RequestMapping(\"/api/auth\").")
    add_spec_table(doc, "3.1.",
        fields=[("authenticationService", "AuthenticationService",
                 "Сервис бизнес-логики аутентификации")],
        methods=[
            ("signUp(request)", "ResponseEntity<JwtResponse>", "public",
             "Регистрация нового пользователя"),
            ("signIn(request)", "ResponseEntity<JwtResponse>", "public",
             "Вход существующего пользователя"),
            ("refresh(request)", "ResponseEntity<JwtResponse>", "public",
             "Обновление токена доступа"),
        ])

    # AuthenticationService spec
    add_para(doc, "Класс AuthenticationService", bold=True, indent=True, space_before=4)
    add_para(doc,
        "Сервисный класс (@Service), реализующий бизнес-логику аутентификации.")
    add_spec_table(doc, "3.2.",
        fields=[
            ("userService",       "UserService",         "Сервис управления пользователями"),
            ("jwtService",        "JwtService",          "Сервис работы с JWT"),
            ("passwordEncoder",   "PasswordEncoder",     "Кодировщик паролей BCrypt"),
            ("authenticationManager", "AuthenticationManager", "Менеджер аутентификации Spring"),
        ],
        methods=[
            ("signUp(request)",    "JwtAuthResponse", "public", "Регистрация, хэширование пароля, выдача токенов"),
            ("signIn(request)",    "JwtAuthResponse", "public", "Проверка учётных данных, выдача токенов"),
            ("refresh(request)",   "JwtAuthResponse", "public", "Обновление токенов по refresh-токену"),
            ("buildResponse(user)","JwtAuthResponse", "private","Формирование ответа с access и refresh токенами"),
        ])

    # JwtService spec
    add_para(doc, "Класс JwtService", bold=True, indent=True, space_before=4)
    add_para(doc,
        "Сервисный класс (@Service), отвечающий за генерацию, валидацию и разбор JWT-токенов.")
    add_spec_table(doc, "3.3.",
        fields=[
            ("jwtSecret",                 "String", "Секретный ключ из application.yaml"),
            ("expirationMs",              "long",   "Время жизни access-токена (900 000 мс)"),
            ("refreshExpirationMs",       "long",   "Время жизни refresh-токена (604 800 000 мс)"),
            ("registrationExpirationMs",  "long",   "Время жизни регистрационного токена (600 000 мс)"),
        ],
        methods=[
            ("generateToken(user)",          "String",  "public", "Генерация access-токена с claims (id, phone)"),
            ("generateRefreshToken(user)",   "String",  "public", "Генерация refresh-токена"),
            ("isTokenValid(token, user)",    "boolean", "public", "Валидация токена по subject и сроку"),
            ("extractSubject(token)",        "String",  "public", "Извлечение subject (phone) из токена"),
            ("generateRegistrationToken()", "String",   "public", "Генерация токена для OAuth-регистрации"),
            ("parseRegistrationToken()",    "Claims",   "public", "Разбор регистрационного токена"),
        ])

    # JwtAuthenticationFilter spec
    add_para(doc, "Класс JwtAuthenticationFilter", bold=True, indent=True, space_before=4)
    add_para(doc,
        "Фильтр Spring Security (OncePerRequestFilter), выполняющий JWT-аутентификацию "
        "для каждого входящего HTTP-запроса.")
    add_spec_table(doc, "3.4.",
        fields=[
            ("jwtService",   "JwtService",   "Сервис работы с токенами"),
            ("userService",  "UserService",  "Сервис загрузки пользователей"),
        ],
        methods=[
            ("doFilterInternal(req, res, chain)", "void", "protected",
             "Извлечение токена из заголовка Authorization, валидация, "
             "установка контекста безопасности"),
        ])

    add_para(doc, "Клиентская часть (Android).", bold=True, indent=True, space_before=6)

    # LoginActivity spec
    add_para(doc, "Класс LoginActivity", bold=True, indent=True, space_before=4)
    add_para(doc,
        "Активность (AppCompatActivity), являющаяся точкой входа в приложение. "
        "Проверяет наличие действующей сессии и перенаправляет пользователя.")
    add_spec_table(doc, "3.5.",
        fields=[],
        methods=[
            ("onCreate(savedInstanceState)", "void",    "protected",
             "Инициализация activity; при наличии refresh-токена — переход в MainActivity"),
            ("hasActiveSession()",           "boolean", "private",
             "Проверка наличия refresh-токена в TokenManager"),
            ("openMainScreen()",             "void",    "private",
             "Запуск MainActivity и завершение LoginActivity"),
        ])

    # LoginViewModel spec
    add_para(doc, "Класс LoginViewModel", bold=True, indent=True, space_before=4)
    add_para(doc,
        "ViewModel архитектуры MVVM, хранящий состояние формы и результаты операций "
        "аутентификации. Взаимодействует с LoginRepository в фоновом потоке.")
    add_spec_table(doc, "3.6.",
        fields=[
            ("loginFormState", "MutableLiveData<LoginFormState>", "Состояние валидации формы"),
            ("loginResult",    "MutableLiveData<LoginResult>",    "Результат операции входа"),
            ("signUpResult",   "MutableLiveData<LoginResult>",    "Результат операции регистрации"),
            ("executor",       "Executor",                        "Однопоточный пул для фоновых задач"),
            ("loginRepository","LoginRepository",                 "Репозиторий аутентификации"),
        ],
        methods=[
            ("login(phone, password)",         "void", "public",
             "Выполняет вход в фоновом потоке, публикует loginResult"),
            ("signUp(firstName,…, password)",  "void", "public",
             "Выполняет регистрацию в фоновом потоке, публикует signUpResult"),
            ("loginDataChanged(phone, pass)",  "void", "public",
             "Валидирует поля формы, публикует loginFormState"),
            ("isPhoneValid(phone)",            "bool", "private", "Проверка номера телефона"),
            ("isPasswordValid(password)",      "bool", "private", "Проверка пароля (>5 символов)"),
        ])

    # LoginRepository spec
    add_para(doc, "Класс LoginRepository", bold=True, indent=True, space_before=4)
    add_para(doc,
        "Класс репозитория, инкапсулирующий логику взаимодействия с REST API "
        "аутентификации и управления токенами.")
    add_spec_table(doc, "3.7.",
        fields=[
            ("authApi",      "AuthApi",      "Retrofit-интерфейс для API аутентификации"),
            ("tokenManager", "TokenManager", "Менеджер хранения токенов"),
        ],
        methods=[
            ("login(phone, password)",               "Result<String>",                   "public", "Вход по паролю, сохранение токенов"),
            ("signUp(firstName,…, password)",         "Result<String>",                   "public", "Регистрация, сохранение токенов"),
            ("getPasskeyRegistrationOptions(phone)",  "Result<PasskeyRegistrationOptions>","public", "Получение параметров регистрации Passkey"),
            ("completePasskeyRegistration(…)",        "Result<String>",                   "public", "Завершение регистрации Passkey"),
            ("startPasskeyAuthentication()",          "Result<PasskeyAssertionOptions>",  "public", "Инициализация аутентификации Passkey"),
            ("finishPasskeyAuthentication(…)",        "Result<String>",                   "public", "Завершение аутентификации Passkey"),
            ("logout()",                              "void",                             "public", "Очистка токенов"),
        ])

    # TokenManager spec
    add_para(doc, "Класс TokenManager", bold=True, indent=True, space_before=4)
    add_para(doc,
        "Класс управления токенами, обеспечивающий безопасное хранение access и refresh "
        "токенов в EncryptedSharedPreferences с шифрованием AES-256-GCM.")
    add_spec_table(doc, "3.8.",
        fields=[
            ("prefs",         "SharedPreferences", "Зашифрованное хранилище AndroidX Security"),
            ("BUFFER_MS",     "long (30 000)",     "Буфер 30 с для предотвращения гонки при обновлении"),
            ("ACCESS_TOKEN_TTL_MS", "long (900 000)", "TTL access-токена (15 мин)"),
        ],
        methods=[
            ("saveTokens(tokens)",         "void",    "public", "Сохранение пары токенов и времени истечения"),
            ("getAccessToken()",           "String",  "public", "Получение access-токена"),
            ("getRefreshToken()",          "String",  "public", "Получение refresh-токена"),
            ("isAccessTokenExpired()",     "boolean", "public", "Проверка истечения с учётом буфера"),
            ("clearTokens()",              "void",    "public", "Удаление всех токенов"),
        ])

    # PasskeyAuthManager spec
    add_para(doc, "Класс PasskeyAuthManager", bold=True, indent=True, space_before=4)
    add_para(doc,
        "Менеджер аутентификации через Passkey, координирующий взаимодействие между "
        "AndroidX CredentialManager и серверным API.")
    add_spec_table(doc, "3.9.",
        fields=[
            ("loginRepository",  "LoginRepository",  "Репозиторий аутентификации"),
            ("credentialManager","CredentialManager", "AndroidX Credential Manager"),
            ("ioExecutor",       "Executor",          "Пул для сетевых операций"),
            ("mainHandler",      "Handler",           "Handler главного потока"),
        ],
        methods=[
            ("authenticate(activity, callback)",        "void", "public",  "Вход через Passkey (аутентификация существующего пользователя)"),
            ("register(activity, phone, callback)",     "void", "public",  "Регистрация нового Passkey"),
            ("requestGetCredential(activity,…)",        "void", "private", "Запрос существующего ключа через CredentialManager"),
            ("finishAuthentication(nonce, json, cb)",   "void", "private", "Финализация аутентификации на сервере"),
            ("requestCreateCredential(activity,…)",     "void", "private", "Создание нового ключа через CredentialManager"),
            ("finishRegistration(nonce, json, cb)",     "void", "private", "Финализация регистрации на сервере"),
        ])

    # AuthInterceptor spec
    add_para(doc, "Класс AuthInterceptor", bold=True, indent=True, space_before=4)
    add_para(doc,
        "OkHttp-перехватчик (Interceptor), автоматически добавляющий Bearer-токен "
        "к запросам и выполняющий обновление токена при получении ответа 401/403.")
    add_spec_table(doc, "3.10.",
        fields=[
            ("tokenManager", "TokenManager", "Менеджер токенов"),
            ("authApi",      "AuthApi",      "Retrofit-интерфейс для обновления токена"),
        ],
        methods=[
            ("intercept(chain)",    "Response", "public",    "Добавление токена, повтор при 401/403"),
            ("getValidToken()",     "String",   "private",   "Получение действующего токена (с обновлением)"),
            ("refreshTokens()",     "String",   "private",   "Синхронное обновление токенов через /api/auth/refresh"),
        ])

    # 3.4
    add_heading(doc, "UML диаграмма классов", level=2, numbered="3.4")

    add_para(doc,
        "В соответствии с рисунком 3.4 представлена UML-диаграмма классов клиентской "
        "части (Android) модуля аутентификации.")

    cl_android = make_class_diagram_android()
    embed_image(doc, cl_android, width_cm=15.5)
    add_figure_caption(doc, "3.4", "UML-диаграмма классов: Android-клиент модуля аутентификации")

    add_para(doc,
        "Центральным классом Android-клиента является LoginViewModel, получающий "
        "запросы от LoginFragment и SignUpCompletionInternalFragment через механизм LiveData. "
        "LoginRepository реализует доступ к сети, делегируя HTTP-запросы интерфейсу AuthApi, "
        "а результаты сохраняя через TokenManager. PasskeyAuthManager координирует "
        "взаимодействие с AndroidX CredentialManager и LoginRepository при Passkey-сценариях. "
        "AuthInterceptor встроен в OkHttpClient и прозрачно управляет жизненным циклом токенов.")

    add_para(doc,
        "В соответствии с рисунком 3.5 представлена UML-диаграмма классов серверной "
        "части модуля аутентификации.")

    cl_backend = make_class_diagram_backend()
    embed_image(doc, cl_backend, width_cm=15.5)
    add_figure_caption(doc, "3.5",
                        "UML-диаграмма классов: серверный модуль аутентификации (Spring Boot)")

    add_para(doc,
        "AuthController принимает HTTP-запросы и делегирует их AuthenticationService. "
        "AuthenticationService использует UserService для работы с репозиторием пользователей "
        "и JwtService для генерации токенов. JwtAuthenticationFilter перехватывает запросы, "
        "валидирует токен через JwtService и устанавливает контекст безопасности "
        "Spring Security. SecurityConfig конфигурирует цепочку фильтров, политику сессий "
        "(STATELESS) и разрешения для эндпоинтов.")

    # 3.5
    add_heading(doc, "Взаимодействие с БД", level=2, numbered="3.5")

    add_para(doc,
        "Взаимодействие с базой данных PostgreSQL осуществляется посредством Spring Data JPA. "
        "Модуль аутентификации работает с двумя таблицами: users и passkey_credentials.")

    add_para(doc, "Сущность User (таблица users).", bold=True, indent=True, space_before=4)
    add_para(doc,
        "Класс User является JPA-сущностью (@Entity, @Table(name = \"users\")), реализующей "
        "интерфейс UserDetails для интеграции со Spring Security. В таблице 3.2 описаны "
        "ключевые поля сущности, используемые в модуле аутентификации.")

    add_table_caption(doc, "3.2", "Поля сущности User, задействованные в аутентификации")
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    for i, h in enumerate(["Поле", "Тип", "Ограничения", "Назначение"]):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            _set_run_font(run, bold=True, size=Pt(11))

    user_fields = [
        ("id",        "UUID",   "@GeneratedValue(UUID)",                  "Первичный ключ"),
        ("username",  "String", "NOT NULL, UNIQUE, 3-32 символа, [a-zA-Z0-9_]", "Отображаемое имя пользователя"),
        ("phone",     "String", "NOT NULL, UNIQUE, паттерн ^\\+?[0-9]{10,15}$", "Логин для Spring Security (getUsername())"),
        ("password",  "String", "NOT NULL",                               "Зашифрованный пароль (BCrypt)"),
        ("firstName", "String", "NOT NULL",                               "Имя"),
        ("lastName",  "String", "NULL",                                   "Фамилия"),
        ("avatarUrl", "String", "NULL",                                   "URL аватара пользователя"),
    ]
    for row_data in user_fields:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            for run in row[i].paragraphs[0].runs:
                _set_run_font(run, size=Pt(10))
    doc.add_paragraph()

    add_para(doc,
        "Особенностью сущности является разделение метода getUsername() (возвращает phone — "
        "логин для Spring Security) и getHandle() (возвращает username — отображаемый "
        "псевдоним). Это обусловлено требованиями интерфейса UserDetails.")

    add_para(doc, "Интерфейс UserRepository.", bold=True, indent=True, space_before=4)
    add_para(doc,
        "JPA-репозиторий (JpaRepository<User, UUID>) предоставляет следующие методы, "
        "используемые в модуле аутентификации:")

    repo_methods = [
        ("findByPhone(phone)", "Поиск пользователя по номеру телефона (для аутентификации)"),
        ("existsByPhone(phone)", "Проверка уникальности телефона при регистрации"),
        ("existsByUsername(username)", "Проверка уникальности имени пользователя при регистрации"),
        ("existsByEmail(email)", "Проверка уникальности email при регистрации"),
    ]
    for method, desc in repo_methods:
        p = doc.add_paragraph()
        run_b = p.add_run(f"{method} — ")
        _set_run_font(run_b, bold=True)
        run_d = p.add_run(desc + ".")
        _set_run_font(run_d)
        _fmt_para(p)

    add_para(doc, "Сущность PasskeyCredential (таблица passkey_credentials).",
             bold=True, indent=True, space_before=4)
    add_para(doc,
        "Сущность хранит публичные ключи пользователей, зарегистрированных через Passkey. "
        "Управление жизненным циклом Passkey-учётных данных осуществляется сервисом "
        "PasskeyService через PasskeyCredentialRepository (JpaRepository<PasskeyCredential, String>). "
        "Первичным ключом является credentialId — уникальный идентификатор ключа, "
        "выданный аутентификатором при регистрации.")

    add_para(doc,
        "Конфигурация базы данных задана в файле application.yaml: "
        "PostgreSQL запускается локально (порт 5432), имя БД — flux. "
        "Для управления схемой используется Hibernate с параметром ddl-auto: update, "
        "что обеспечивает автоматическое применение изменений схемы без потери данных.")

    # 3.6
    add_heading(doc, "Тестирование", level=2, numbered="3.6")

    add_para(doc,
        "Тестирование модуля аутентификации проводилось на уровне модульных тестов "
        "серверной части с использованием фреймворка JUnit 5 и библиотеки AssertJ. "
        "Тестируемым классом является JwtService.")

    add_para(doc, "В таблице 3.3 приведены реализованные тест-кейсы.", space_before=4)

    add_table_caption(doc, "3.3", "Тест-кейсы класса JwtServiceTest")
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    for i, h in enumerate(["№", "Название теста", "Проверяемый сценарий"]):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            _set_run_font(run, bold=True, size=Pt(11))

    tests = [
        ("1", "generateAndExtractAccessToken",
         "Генерация access-токена и корректное извлечение subject (phone)"),
        ("2", "generateRefreshTokenContainsSubject",
         "Refresh-токен содержит корректный subject пользователя"),
        ("3", "registrationTokenHasPurposeClaim",
         "Регистрационный токен содержит claim purpose=oauth-register"),
        ("4", "parseRegistrationTokenRejectsAccessToken",
         "Попытка разобрать access-токен как регистрационный вызывает IllegalArgumentException"),
        ("5", "parseRegistrationTokenRejectsExpiredToken",
         "Истёкший регистрационный токен вызывает RegistrationTokenExpiredException"),
        ("6", "parseRegistrationTokenRejectsMalformed",
         "Некорректный токен вызывает IllegalArgumentException"),
    ]
    for row_data in tests:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            for run in row[i].paragraphs[0].runs:
                _set_run_font(run, size=Pt(10))
    doc.add_paragraph()

    add_para(doc,
        "Все шесть тест-кейсов успешно пройдены. Конфигурация тестов использует "
        "ReflectionTestUtils для внедрения значений jwtSecret, expirationMs и "
        "refreshExpirationMs без запуска полного контекста Spring. "
        "Дополнительно выполнено ручное тестирование потока аутентификации: "
        "отправка запросов к эндпоинтам /api/auth/sign-up, /api/auth/sign-in и "
        "/api/auth/refresh с проверкой корректности токенов, защищённых эндпоинтов "
        "(возврат 401 при отсутствии или некорректном токене) и логики автоматического "
        "обновления в AuthInterceptor при получении ответа 401.")

    # 3.7
    add_heading(doc, "Руководство пользователя", level=2, numbered="3.7")

    add_para(doc,
        "Данный раздел описывает порядок работы пользователя с экранами модуля "
        "аутентификации мобильного мессенджера Flux Messenger.")

    add_para(doc, "Запуск приложения.", bold=True, indent=True, space_before=4)
    add_para(doc,
        "При первом запуске приложения открывается LoginActivity. Если на устройстве "
        "уже сохранён действующий refresh-токен (пользователь ранее вошёл в систему), "
        "приложение автоматически переходит в MainActivity без отображения экранов входа.")

    add_para(doc, "Экран приветствия (WelcomeAuthFragment).", bold=True, indent=True, space_before=4)
    add_para(doc,
        "На экране приветствия пользователю предлагается три варианта действий:")
    welcome_steps = [
        "нажать кнопку «Войти» для перехода к форме входа по паролю;",
        "нажать кнопку «Зарегистрироваться» для перехода к форме регистрации;",
        "нажать кнопку «Войти через Passkey» для аутентификации биометрией — если Passkey "
        "не зарегистрирован, приложение предложит его зарегистрировать.",
    ]
    for s in welcome_steps:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(s)
        _set_run_font(run)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    add_para(doc, "Вход в систему (LoginFragment).", bold=True, indent=True, space_before=4)
    add_para(doc,
        "Для входа необходимо ввести номер телефона в поле «Телефон» и пароль в поле «Пароль». "
        "Кнопка «Войти» становится доступной только после ввода корректных данных (телефон — "
        "непустой, пароль — более 5 символов). При нажатии кнопки выполняется запрос к серверу; "
        "в случае успеха пользователь переходит в главный экран мессенджера. "
        "При ошибке аутентификации отображается уведомление (Toast).")

    add_para(doc, "Регистрация (SignUpAuthFragment → SignUpCompletionInternalFragment).",
             bold=True, indent=True, space_before=4)
    add_para(doc,
        "Регистрация выполняется в два шага. На первом шаге необходимо ввести номер "
        "телефона (не менее 11 цифр), пароль (более 5 символов) и подтверждение пароля. "
        "Кнопка «Далее» активируется только при совпадении паролей и корректности всех полей. "
        "На втором шаге вводятся имя (обязательно), фамилия (опционально) и имя пользователя "
        "(не менее 3 символов, только латинские буквы, цифры и «_»). После нажатия кнопки "
        "«Готово» выполняется регистрация на сервере и автоматический вход в приложение.")

    add_para(doc, "Регистрация и вход через Passkey (SignUpCompletion3rdPartyFragment).",
             bold=True, indent=True, space_before=4)
    add_para(doc,
        "Для регистрации Passkey необходимо заполнить имя, имя пользователя и номер телефона, "
        "затем нажать кнопку «Войти». Приложение запросит подтверждение биометрией или PIN-кодом. "
        "После подтверждения ключ регистрируется на сервере и пользователь переходит в "
        "главный экран. При повторном входе через Passkey достаточно нажать кнопку "
        "«Войти через Passkey» на экране приветствия и подтвердить биометрией.")

    page_break(doc)

    # ── ЗАКЛЮЧЕНИЕ ────────────────────────────────────────────────────────────
    add_heading(doc, "Заключение", level=1)

    add_para(doc,
        "В ходе выполнения курсовой работы был разработан модуль аутентификации и "
        "авторизации для мобильного мессенджера Flux Messenger. Реализованы серверная "
        "часть на основе Spring Boot с поддержкой JWT-токенов и стандарта WebAuthn/Passkey, "
        "а также клиентская часть Android-приложения с архитектурой MVVM.")

    add_para(doc,
        "В результате выполнения курсовой работы были решены все поставленные задачи:")

    results = [
        "проведён анализ предметной области и современных методов аутентификации в "
        "мобильных приложениях;",
        "спроектирована архитектура модуля в соответствии с паттерном MVVM и принципами "
        "разделения ответственности;",
        "реализован REST API аутентификации с эндпоинтами для входа, регистрации "
        "и обновления токенов;",
        "реализована поддержка беспарольного входа посредством WebAuthn/Passkey "
        "с использованием AndroidX Credentials API и Spring Security WebAuthn;",
        "обеспечено безопасное хранение токенов на устройстве посредством "
        "EncryptedSharedPreferences с шифрованием AES-256-GCM;",
        "реализовано автоматическое обновление токенов через OkHttp-перехватчик "
        "без участия пользователя;",
        "проведено модульное тестирование JwtService с покрытием шести тест-кейсов.",
    ]
    for r in results:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(r)
        _set_run_font(run)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    add_para(doc,
        "Разработанный модуль обеспечивает надёжную защиту учётных данных пользователей "
        "и удобство использования приложения. Применение технологии Passkey позволяет "
        "полностью отказаться от паролей, снизив риски фишинга и утечки данных. "
        "В дальнейшем возможно расширение модуля добавлением двухфакторной аутентификации "
        "и поддержки OAuth 2.0 для входа через сторонние сервисы.")

    page_break(doc)

    # ── СПИСОК ЛИТЕРАТУРЫ ─────────────────────────────────────────────────────
    add_heading(doc, "Список литературы", level=1)

    references = [
        "Spring Security Reference Documentation. — URL: "
        "https://docs.spring.io/spring-security/reference/ (дата обращения: 18.05.2025).",
        "JSON Web Token (JWT). RFC 7519. — URL: https://www.rfc-editor.org/rfc/rfc7519 "
        "(дата обращения: 18.05.2025).",
        "Web Authentication: An API for accessing Public Key Credentials Level 3. W3C. — "
        "URL: https://www.w3.org/TR/webauthn-3/ (дата обращения: 18.05.2025).",
        "Android Developers. Credential Manager. — URL: "
        "https://developer.android.com/identity/sign-in/credential-manager "
        "(дата обращения: 18.05.2025).",
        "Android Developers. EncryptedSharedPreferences. — URL: "
        "https://developer.android.com/reference/androidx/security/crypto/EncryptedSharedPreferences "
        "(дата обращения: 18.05.2025).",
        "Square Open Source. Retrofit. — URL: https://square.github.io/retrofit/ "
        "(дата обращения: 18.05.2025).",
        "Square Open Source. OkHttp. — URL: https://square.github.io/okhttp/ "
        "(дата обращения: 18.05.2025).",
        "JJWT — Java JWT Library. GitHub. — URL: https://github.com/jwtk/jjwt "
        "(дата обращения: 18.05.2025).",
        "Android Developers. Guide to app architecture. — URL: "
        "https://developer.android.com/topic/architecture "
        "(дата обращения: 18.05.2025).",
        "СТО 005–2020. Система менеджмента качества. Учебно-методическая деятельность. "
        "Оформление курсовых проектов (работ) и выпускных квалификационных работ "
        "технических направлений подготовки и специальностей. — Иркутск: ИРНИТУ, 2020. — 40 с.",
    ]
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {ref}")
        _set_run_font(run)
        pf = p.paragraph_format
        pf.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = PARA_INDENT
        pf.space_before      = Pt(2)
        pf.space_after       = Pt(2)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    return doc


if __name__ == "__main__":
    print("Generating diagrams and building document…")
    doc = build_report()
    out = "/home/user/Flux-Messenger/Отчёт_Flux_Messenger_Аутентификация.docx"
    doc.save(out)
    print(f"Saved: {out}")
